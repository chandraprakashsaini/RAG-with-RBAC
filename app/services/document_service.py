from __future__ import annotations

import asyncio
from pathlib import Path
from typing import BinaryIO, Optional
from uuid import UUID, uuid4

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.core.config import get_settings
from app.db.models import Document, DocumentPermission, UserRole
from app.models.chunking import ChunkingConfig
from app.services.chunking_service import build_chunks_payload
from app.services.vector_store_service import (
    delete_document_chunks,
    store_document_chunks,
)

SUPPORTED_MIME_TYPES: dict[str, str] = {
    ".txt": "text/plain",
    ".md": "text/markdown",
    ".csv": "text/csv",
    ".json": "application/json",
    ".xml": "application/xml",
    ".yml": "application/x-yaml",
    ".yaml": "application/x-yaml",
    ".html": "text/html",
    ".htm": "text/html",
}

try:
    import fitz  # noqa: F401

    SUPPORTED_MIME_TYPES[".pdf"] = "application/pdf"
except ImportError:
    pass

try:
    import docx  # noqa: F401

    SUPPORTED_MIME_TYPES[".docx"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
except ImportError:
    pass


def detect_mime(filename: str) -> str | None:
    ext = Path(filename).suffix.lower()
    return SUPPORTED_MIME_TYPES.get(ext)


def _read_text_content_sync(file_path: Path, mime_type: str) -> str:
    if mime_type == "application/pdf":
        import fitz

        doc = fitz.open(str(file_path))
        text = "\n".join(page.get_text() for page in doc)
        doc.close()
        return text

    if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        import docx

        doc = docx.Document(str(file_path))
        return "\n".join(p.text for p in doc.paragraphs)

    with open(file_path, encoding="utf-8", errors="replace") as f:
        return f.read()


def read_text_content(file_path: Path, mime_type: str) -> str:
    return _read_text_content_sync(file_path, mime_type)


async def read_text_content_async(file_path: Path, mime_type: str) -> str:
    return await asyncio.to_thread(_read_text_content_sync, file_path, mime_type)


def _save_upload_file_sync(
    file_contents: BinaryIO, dest: Path, max_bytes: int
) -> int:
    """Stream the upload to disk in chunks, enforcing a size cap.

    Returns the number of bytes written. Raises ValueError if the cap is exceeded.
    """
    written = 0
    chunk_size = 1024 * 1024  # 1 MiB
    with open(dest, "wb") as f:
        while True:
            chunk = file_contents.read(chunk_size)
            if not chunk:
                break
            written += len(chunk)
            if written > max_bytes:
                f.close()
                try:
                    dest.unlink()
                except OSError:
                    pass
                raise ValueError(f"Upload exceeds maximum size of {max_bytes} bytes")
            f.write(chunk)
    return written


async def save_upload_file(
    file_contents: BinaryIO,
    original_filename: str,
    max_bytes: int = 25_000_000,
) -> tuple[str, str, int]:
    settings = get_settings()
    docs_dir = settings.documents_dir
    docs_dir.mkdir(parents=True, exist_ok=True)

    doc_uuid = uuid4()
    ext = Path(original_filename).suffix
    safe_filename = f"{doc_uuid}{ext}"
    file_path = docs_dir / safe_filename

    file_size = await asyncio.to_thread(
        _save_upload_file_sync, file_contents, file_path, max_bytes
    )

    return str(doc_uuid), str(file_path), file_size


async def process_uploaded_text(
    text: str,
    document_name: str | None = None,
    chunking: ChunkingConfig | None = None,
) -> tuple[str, int]:
    if chunking is None:
        chunking = ChunkingConfig()

    chunking_result = await build_chunks_payload(text, chunking)
    doc_uuid = uuid4()
    stored_ids = await store_document_chunks(
        document_id=str(doc_uuid),
        chunking_result=chunking_result,
        document_name=document_name,
    )
    return str(doc_uuid), len(stored_ids)


async def create_document_record(
    db: AsyncSession,
    user_id: UUID,
    original_filename: str,
    file_path: str,
    file_size: int,
    mime_type: str,
    document_id: UUID,
    chunk_count: int,
) -> Document:
    doc = Document(
        user_id=user_id,
        original_filename=original_filename,
        file_path=file_path,
        file_size=file_size,
        mime_type=mime_type,
        document_id=document_id,
        chunk_count=chunk_count,
    )
    db.add(doc)
    await db.commit()
    await db.refresh(doc)
    return doc


async def list_documents(
    db: AsyncSession, user_id: UUID, role_name: str, limit: int = 100
) -> list[Document]:
    owned_result = await db.execute(
        select(Document)
        .where(Document.user_id == user_id)
        .options(selectinload(Document.user))
        .order_by(Document.created_at.desc())
        .limit(limit)
    )
    owned = owned_result.scalars().all()

    shared_result = await db.execute(
        select(Document)
        .join(DocumentPermission, Document.id == DocumentPermission.document_id)
        .join(UserRole, DocumentPermission.role_id == UserRole.id)
        .where(UserRole.name == role_name)
        .options(selectinload(Document.user))
        .order_by(Document.created_at.desc())
        .limit(limit)
    )
    shared = shared_result.scalars().all()

    seen = set()
    docs = []
    for doc in owned + shared:
        if doc.id not in seen:
            seen.add(doc.id)
            docs.append(doc)
    docs.sort(key=lambda d: d.created_at, reverse=True)
    return docs[:limit]


async def get_user_document(
    db: AsyncSession, doc_id: UUID, user_id: UUID, role_name: str
) -> Document | None:
    result = await db.execute(
        select(Document).where(Document.id == doc_id, Document.user_id == user_id).options(selectinload(Document.user))
    )
    doc = result.scalar_one_or_none()
    if doc:
        return doc

    result = await db.execute(
        select(Document)
        .join(DocumentPermission, Document.id == DocumentPermission.document_id)
        .join(UserRole, DocumentPermission.role_id == UserRole.id)
        .where(
            Document.id == doc_id,
            UserRole.name == role_name,
            DocumentPermission.can_read == True,
        )
        .options(selectinload(Document.user))
    )
    return result.scalar_one_or_none()


async def check_document_delete_permission(
    db: AsyncSession, doc_id: UUID, user_id: UUID, role_name: str
) -> Document | None:
    result = await db.execute(
        select(Document).where(Document.id == doc_id, Document.user_id == user_id)
    )
    doc = result.scalar_one_or_none()
    if doc:
        return doc

    result = await db.execute(
        select(Document)
        .join(DocumentPermission, Document.id == DocumentPermission.document_id)
        .join(UserRole, DocumentPermission.role_id == UserRole.id)
        .where(
            Document.id == doc_id,
            UserRole.name == role_name,
            DocumentPermission.can_delete == True,
        )
    )
    return result.scalar_one_or_none()


async def delete_document(db: AsyncSession, doc: Document) -> None:
    doc_file = Path(doc.file_path)
    if doc_file.exists():
        doc_file.unlink()

    await delete_document_chunks(str(doc.document_id))

    await db.delete(doc)
    await db.commit()


# ── Document permission CRUD ──────────────────────────────────────────


async def list_document_permissions(
    db: AsyncSession, document_id: UUID
) -> list[DocumentPermission]:
    result = await db.execute(
        select(DocumentPermission)
        .where(DocumentPermission.document_id == document_id)
        .options(selectinload(DocumentPermission.role))
        .options(selectinload(DocumentPermission.grantor))
    )
    return list(result.scalars().all())


async def grant_document_permission(
    db: AsyncSession,
    document_id: UUID,
    role_id: UUID,
    can_read: bool,
    can_write: bool,
    can_delete: bool,
    granted_by: UUID,
) -> DocumentPermission:
    existing = await db.execute(
        select(DocumentPermission).where(
            DocumentPermission.document_id == document_id,
            DocumentPermission.role_id == role_id,
        )
    )
    if existing.scalar_one_or_none():
        raise ValueError("This role already has permissions on this document")

    perm = DocumentPermission(
        document_id=document_id,
        role_id=role_id,
        can_read=can_read,
        can_write=can_write,
        can_delete=can_delete,
        granted_by=granted_by,
    )
    db.add(perm)
    await db.commit()
    await db.refresh(perm)

    result = await db.execute(
        select(DocumentPermission)
        .where(DocumentPermission.id == perm.id)
        .options(selectinload(DocumentPermission.role))
        .options(selectinload(DocumentPermission.grantor))
    )
    return result.scalar_one()


async def update_document_permission(
    db: AsyncSession,
    permission_id: UUID,
    can_read: Optional[bool],
    can_write: Optional[bool],
    can_delete: Optional[bool],
) -> DocumentPermission | None:
    result = await db.execute(
        select(DocumentPermission).where(DocumentPermission.id == permission_id)
    )
    perm = result.scalar_one_or_none()
    if not perm:
        return None

    if can_read is not None:
        perm.can_read = can_read
    if can_write is not None:
        perm.can_write = can_write
    if can_delete is not None:
        perm.can_delete = can_delete

    await db.commit()
    await db.refresh(perm)

    result = await db.execute(
        select(DocumentPermission)
        .where(DocumentPermission.id == perm.id)
        .options(selectinload(DocumentPermission.role))
        .options(selectinload(DocumentPermission.grantor))
    )
    return result.scalar_one()


async def revoke_document_permission(
    db: AsyncSession, permission_id: UUID
) -> bool:
    result = await db.execute(
        select(DocumentPermission).where(DocumentPermission.id == permission_id)
    )
    perm = result.scalar_one_or_none()
    if not perm:
        return False

    await db.delete(perm)
    await db.commit()
    return True
